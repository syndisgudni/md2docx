#! /usr/bin/python

# MD > HTML
import mistletoe
# HTML parser
from bs4 import BeautifulSoup as bs
from bs4.element import Tag, NavigableString
from prism import highlight
# DOCX
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.text.run import Run
# Utils
import argparse
import glob
import os
from PIL import Image, ImageOps
from colour import Color

DOC_TEXT_WIDTH = "450pt"
COLOR_SYN_BLUE = "#057d9f"

def font_size(sz):
    return Pt(float(sz.split('pt')[0]))

def font_color(c):
    color = Color(c.strip())
    rgb = color.rgb
    return RGBColor(int(rgb[0]*255), int(rgb[1]*255), int(rgb[2]*255))

def par_align(j):
    match j:
        case 'justify':
            return WD_ALIGN_PARAGRAPH.JUSTIFY
        case 'center':
            return WD_ALIGN_PARAGRAPH.CENTER
        case 'left':
            return WD_ALIGN_PARAGRAPH.LEFT
        case 'right':
            return WD_ALIGN_PARAGRAPH.RIGHT

def set_image_border(img, sz, color):
    i = img._inline
    pic = i.graphic.graphicData.pic
    spPr = pic.spPr

    ln = OxmlElement('a:ln', {'w': '12700'})
    spPr.append(ln)

    solidFill = OxmlElement('a:solidFill')
    ln.append(solidFill)
    srgbClr = OxmlElement('a:srgbClr', {'val': '666666'})
    solidFill.append(srgbClr)
    
    prstDash = OxmlElement('a:prstDash', {'val': 'solid'})
    ln.append(prstDash)

def set_paragraph_border(par, **kwargs):
    """
    Set paragraph`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    p = par._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    
    # list over all available tags
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            if 'color' in edge_data:
                color = edge_data['color'][1:]
                edge_data['color'] = color
            if 'sz' in edge_data:
                sz = int(edge_data['sz'][:-2])
                edge_data['sz'] = str(sz*2)

            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = pBdr.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                pBdr.append(element)
            # looks like order of attributes is important
            for key in ["color", "space", "sz", "val", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def set_cell_border(c, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = c._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existance, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            if 'color' in edge_data:
                color = edge_data['color'][1:]
                edge_data['color'] = color
            if 'sz' in edge_data:
                sz = int(edge_data['sz'][:-2])
                edge_data['sz'] = str(sz*2)

            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            # looks like order of attributes is important
            for key in ["color", "space", "sz", "val", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def set_cell_bg_color(c, color):
    color = color[1:]
    tblCell = c._tc
    tblCellProperties = tblCell.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    tblCellProperties.append(shading)

def dict_to_style(dict):
    return ';'.join(['%s:%s' % (key, value) for (key, value) in dict.items()])

def style_to_dict(style):
    d = {}
    for s in style.split(';'):
        if s =='': continue
        k,v = s.split(':')
        d[k] = v
    return d

def get_style(tag):
    style = {}
    tags = []
    # Build tag stack
    while tag.parent:
        tags.append(tag)
        tag = tag.parent
    # Traverse tag stack and build style object
    while len(tags) > 0:
        tag = tags.pop()
        if not tag.name is None and 'style' in tag.attrs:
            style.update(style_to_dict(tag['style']))
    return style

def table_dimensions(table):
    rows = len(table.find_all('tr'))
    cols = len(table.tr.select('th,td'))
    return rows, cols

def generate_reqres(contents):
    table = '''<table class="reqres">
            <thead>
            <tr>
                <th style="" colspan=2><b>Reproduction example</b></th><th style=""></th>
            </tr>
            </thead>
            <tbody>{}</tbody>
        </table>'''
    rows = ''
    for key, val in [s.split('\n', 1) for s in contents.split('--- ')[1:]]:
        blurb =''
        if '|' in key: key, blurb = key.split('|', 1)
        row = '''<tr>
                <td style="width:50pt"><b>{}</b><i style="font-size:10pt">{}</i></td>
                <td style="width:400pt">
                    <pre><code class="language-http">{}</code></pre>
                </td>
            </tr>'''.format(key.strip(),'\n'+blurb.strip(),highlight(val.strip(),'http'))
        rows += row
    return table.format(rows)

def apply_html_style(soup):
    # Global
    base_style = {
        'font-family': 'Calibri',
        'font-size': '11pt'
    }

    # Parse custom reproduction example code block
    # to HTML table for later styling
    for t in soup.select('pre > code.language-reqres'):
        inner_html = generate_reqres(t.get_text())
        new_soup = bs(inner_html, 'html.parser')
        t.parent.replace_with(new_soup)

    # Text
    for t in soup.select('p, th, td, li'):
        style = base_style.copy()
        style['text-align'] = 'justify'
        t['style'] = dict_to_style(style)

    # Inline Code
    for t in soup.find_all('code'):
        style = base_style.copy()
        style['font-family'] = 'Roboto Mono'
        style['font-size'] = '9pt'
        t['style'] = dict_to_style(style)

    # Hyperlinks
    for t in soup.find_all('a'):
        style = base_style.copy()
        style['color'] = '#1155cc'
        style['text-decoration-line'] = 'underline'
        t['style'] = dict_to_style(style)

    # Code blocks
    for t in soup.select('pre > code'):
        style = base_style.copy()
        style['font-family'] = 'Roboto Mono'
        style['font-size'] = '9pt'
        style['text-align'] = 'left'
        style['border'] = '1px solid #000000'
        t['style'] = dict_to_style(style)

        # Apply Prism syntax highlighting
        lang = ''
        if 'class' in t.attrs:
            lang = t['class'][0].removeprefix('language-')
        code = t.get_text()
        t.clear()
        if lang == '':
            inner_html = code
        else:
            inner_html = highlight(code, lang)
        new_soup = bs(inner_html, 'html.parser')
        t.append(new_soup)

    # Image
    for i in soup.find_all('img'):
        style = base_style.copy()
        style['width'] = DOC_TEXT_WIDTH
        style['border'] = '1px solid #000000'
        i['style'] = dict_to_style(style)

    # Captions
    for p in soup.select('p:has(img), pre + p, table + p'):
        style = base_style.copy()
        style['font-size'] = '10pt'
        style['text-align'] = 'center'
        p['style'] = dict_to_style(style)
        for t in p.find_all('code'):
            style = base_style.copy()
            if 'style' in t.attrs:
                style = style_to_dict(t['style'])
            style['font-size'] = '8pt'
            t['style'] = dict_to_style(style)


    # Headings
    heading_style = base_style.copy()
    heading_style['color'] = '#666666'
    heading_style['text-align'] = 'left'

    for t in soup.find_all('h1'):
        style = heading_style.copy()
        style['font-size'] = '20pt'
        style['font-weight'] = 'bold'
        t['style'] = dict_to_style(style)

    for t in soup.find_all('h2'):
        style = heading_style.copy()
        style['font-size'] = '18pt'
        style['font-weight'] = 'bold'
        t['style'] = dict_to_style(style)
        
    for t in soup.find_all('h3'):
        style = heading_style.copy()
        style['font-size'] = '14pt'
        style['font-weight'] = 'bold'
        style['color'] = '#434343'
        t['style'] = dict_to_style(style)

    for t in soup.find_all('h4'):
        style = heading_style.copy()
        style['font-size'] = '12pt'
        t['style'] = dict_to_style(style)

    # Tables
    table_style = base_style.copy()
    table_style['border-collapse'] = 'collapse'

    ## Remove empty table heads
    for t in soup.find_all('table'):
        header_empty = True
        if not t.thead:
            break
        for th in t.thead.tr:
            if not th.text.rstrip() == '':
                header_empty = False
                break
        if header_empty:
            t.thead.decompose()
    ## Remove those weird alignment attributes on every table cell
    for t in soup.select('tr > *'):
        if 'align' in t.attrs: del t['align']
        ## Force left-alignment on table cells
        style = style_to_dict(t['style'])
        style['text-align'] = 'left'
        t['style'] = dict_to_style(style)
    ## All tables
    for t in soup.find_all('table'):
        t['style'] = dict_to_style(table_style)
    ## Summary tables
    for t in soup.select("h2 + table"):
        ### Top row
        for c in t.tr.find_all('td'):
            if 'style' in c.attrs: style = style_to_dict(c['style'])
            else: style = base_style.copy()
            style['border-top'] = '6px solid %s' % COLOR_SYN_BLUE
            c['style'] = dict_to_style(style)
        for r in t.tbody.find_all('tr'):
            ### First column
            c = r.td
            if 'style' in c.attrs: style = style_to_dict(c['style'])
            else: style = base_style.copy()
            style['text-align'] = 'right'
            style['color'] = '#ffffff'
            style['background-color'] = COLOR_SYN_BLUE
            style['width'] = '82.512pt'
            style['border-top'] = '6px solid %s' % COLOR_SYN_BLUE
            c['style'] = dict_to_style(style)
            ### Second column
            c = r.select('td:nth-child(2)')[0]
            if 'style' in c.attrs: style = style_to_dict(c['style'])
            else: style = base_style.copy()
            style['width'] = '367.488pt'
            c['style'] = dict_to_style(style)
    ## Other tables
    for t in soup.select("*:not(h2) + table"):
        ### Top row
        if t.thead:
            for c in t.thead.find_all('th'):
                if 'style' in c.attrs: style = style_to_dict(c['style'])
                else: style = base_style.copy()
                style['color'] = '#ffffff'
                style['background-color'] = COLOR_SYN_BLUE
                c['style'] = dict_to_style(style)
        ### Every other row
        for r in t.tbody.select('tr:nth-child(even)'):
            for c in r.find_all('td'):
                if 'style' in c.attrs: style = style_to_dict(c['style'])
                else: style = base_style.copy()
                style['background-color'] = '#efefef'
                c['style'] = dict_to_style(style)
    ## Request/response tables
    for t in soup.select(".reqres"):
        c = t.th
        if 'style' in c.attrs: style = style_to_dict(c['style'])
        else: style = base_style.copy()
        style['font-size'] = '12pt'
        c['style'] = dict_to_style(style)
        for r in t.tbody.find_all('tr'):
            ### First column
            c = r.td
            if 'style' in c.attrs: style = style_to_dict(c['style'])
            else: style = base_style.copy()
            style['width'] = '80pt'
            style['font-size'] = '12pt'
            style['color'] = '#434343'
            c['style'] = dict_to_style(style)
            ### Second column
            c = r.select('td:nth-child(2)')[0]
            if 'style' in c.attrs: style = style_to_dict(c['style'])
            else: style = base_style.copy()
            style['width'] = '370pt'
            c['style'] = dict_to_style(style)
        ### Every cell
        for c in t.tbody.find_all('td'):
            if 'style' in c.attrs: style = style_to_dict(c['style'])
            else: style = base_style.copy()
            style['border-bottom'] = '2px solid %s' % COLOR_SYN_BLUE
            c['style'] = dict_to_style(style)
        ### Every other row
        for r in t.tbody.select('tr:nth-child(even)'):
            for c in r.find_all('td'):
                if 'style' in c.attrs: style = style_to_dict(c['style'])
                else: style = base_style.copy()
                del style['background-color']
                c['style'] = dict_to_style(style)
        


class HtmlToDocx:
    def __init__(self, soup: bs):
        self.soup = soup
        self.doc = self.setup_document()
        self.stack = []
    
    def setup_document(self):
        doc = Document()

        # Page setup
        section = doc.sections[-1]
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.59)
        section.bottom_margin = Inches(0.59)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        p_format = doc.styles['Normal'].paragraph_format
        p_format.line_spacing = 1.0
        p_format.space_before = Pt(0)
        p_format.space_after = Pt(11)

        p_format = doc.styles['Heading 1'].paragraph_format
        p_format.line_spacing = 1.0
        p_format.space_before = Pt(0)
        p_format.space_after = Pt(6)

        p_format = doc.styles['Heading 2'].paragraph_format
        p_format.line_spacing = 1.0
        p_format.space_before = Pt(18)
        p_format.space_after = Pt(6)

        p_format = doc.styles['Heading 3'].paragraph_format
        p_format.line_spacing = 1.0
        p_format.space_before = Pt(16)
        p_format.space_after = Pt(4)

        p_format = doc.styles['Heading 4'].paragraph_format
        p_format.line_spacing = 1.0
        p_format.space_before = Pt(14)
        p_format.space_after = Pt(4)

        return doc

    def strip_paragraph(self, p):
        # Clear leading empty runs
        for r in p.runs:
            if not r.text == '\n':
                break
            r.text = ''
        # Clear trailing empty runs
        for r in reversed(p.runs):
            if not r.text == '\n':
                break
            r.text = ''

    def apply_inline_style(self, r, style):
        if 'font-family' in style:
            r.font.name = style['font-family']
        if 'font-size' in style:
            r.font.size = font_size(style['font-size'])
        if 'font-weight' in style:
            r.bold = (style['font-weight'] == 'bold')
        if 'color' in style:
            r.font.color.rgb = font_color(style['color'])
        if 'text-decoration-line' in style and style['text-decoration-line'] == 'underline':
            r.underline = WD_UNDERLINE.SINGLE
        scope = [t.name for t in self.stack]
        r.bold = r.bold or 'strong' in scope or 'b' in scope
        r.italic = r.italic or 'em' in scope or 'i' in scope

    def apply_block_style(self, p, style):
        if 'text-align' in style:
            p.alignment = par_align(style['text-align'])
        if 'border' in style:
            sz,_,color = style['border'].split(' ')
            set_paragraph_border(p,
                top={'sz': sz, 'val': 'single', 'color': color, 'space': '4'},
                bottom={'sz': sz, 'val': 'single', 'color': color, 'space': '4'},
                left={'sz': sz, 'val': 'single', 'color': color, 'space': '4'},
                right={'sz': sz, 'val': 'single', 'color': color, 'space': '4'}
            )

    def apply_cell_style(self, c, style):
        if 'background-color' in style:
            set_cell_bg_color(c, style['background-color'])
        if 'border-top' in style:
            sz,_,color = style['border-top'].split(' ')
            set_cell_border(c, top={'sz': sz, 'val': 'single', 'color': color, 'space': '0'})
        if 'border-bottom' in style:
            sz,_,color = style['border-bottom'].split(' ')
            set_cell_border(c, bottom={'sz': sz, 'val': 'single', 'color': color, 'space': '0'})

    def apply_column_style(self, c, style):
        if 'width' in style:
            c.width = font_size(style['width'])
    
    def apply_table_style(self, t, style):
        t.autofit = True

    def apply_image_style(self, i, style):
        if 'width' in style:
            i.width = font_size(style['width'])
        if 'border' in style:
            sz,_,color = style['border'].split(' ')
            set_image_border(i, sz, color)

    def render_text(self, tag, p):
        if not isinstance(tag, Tag):
            r = p.add_run(tag.text)
            self.apply_inline_style(r, get_style(tag))
            return
        self.stack.append(tag)
        for t in tag:
            self.render_text(t, p)
        self.stack.pop()
        return tag.text

    def render_link(self, tag, p):
        url = ''
        if 'href' in tag.attrs:
            url = tag['href']
        text = tag.text
        
        r_id = p.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        h = OxmlElement('w:hyperlink')
        h.set(qn('r:id'), r_id)

        r = Run(OxmlElement('w:r'), p)
        r.text = text

        self.apply_inline_style(r, get_style(tag))

        h.append(r._element)
        p._p.append(h)

    def render_paragraph(self, tag, p=None):
        # This is a weird workaround because python-docx is bad with pictures
        if p == None:
            if 'img' in [t.name for t in tag.children]:
                self.render_image(tag.img)
                p = self.doc.paragraphs[-1]
            else:
                p = self.doc.add_paragraph()
        self.apply_block_style(p, get_style(tag))
        for t in tag:
            if t.name == 'img': continue
            elif t.name == 'a': self.render_link(t, p)
            else: self.render_text(t, p)
        return p

    def render_table(self, tag):
        rows, cols = table_dimensions(tag)
        t = self.doc.add_table(rows=rows, cols=cols)
        self.apply_table_style(t, get_style(tag))

        soup_rows = tag.find_all('tr')
        for r, row in enumerate(t.rows):
            soup_row = soup_rows[r]
            soup_cells = soup_row.select('th,td')
            for c, cell in enumerate(row.cells):
                soup_cell = soup_cells[c]
                self.render_paragraph(soup_cell, cell.paragraphs[0])
                self.apply_cell_style(cell, get_style(soup_cell))
                par_form = cell.paragraphs[0].paragraph_format
                par_form.space_before = Pt(5)
                par_form.space_after = Pt(5)
        
        for c, col in enumerate(t.columns):
            soup_cell = tag.select('td:nth-child(%d)' % (c+1))[0]
            self.apply_column_style(col, get_style(soup_cell))
                
        # Merge top-row cells in request/response tables
        if 'class' in tag.attrs and 'reqres' in tag['class']:
            t.cell(0, 0).merge(t.cell(0, 1))
            for c in t.columns[1].cells:
                self.strip_paragraph(c.paragraphs[0])

        return t

    def render_heading(self, tag):
        level = int(tag.name[1])
        h = self.doc.add_heading('', level)
        self.apply_block_style(h, get_style(tag))
        par_form = h.paragraph_format
        par_form.line_spacing = 1.0
        par_form.space_before = Pt(16)
        par_form.space_after = Pt(4)
        self.render_text(tag, h)
        return h

    def render_image(self, tag):
        i = self.doc.add_picture(tag['src'], width=font_size(DOC_TEXT_WIDTH))
        self.apply_image_style(i, get_style(tag))
        return i

    def render_code(self, tag):
        # Remove trailing newline
        if len(tag.code.contents) > 1:
            tag.code.contents[-1].replace_with('')
        self.render_paragraph(tag.code)

    def render_list(self, tag):
        if tag.name == 'ul': style = 'List Bullet'
        else: style = 'List Number'
        for i in tag.select('li'):
            self.render_paragraph(i).style = self.doc.styles[style]

    def render_tag(self, tag, level=0):
        self.stack.append(tag)

        rendered = True

        # Parse block-level tags
        match tag.name:
            case 'h1' | 'h2' | 'h3' | 'h4':
                self.render_heading(tag)
            case 'p':
                self.render_paragraph(tag)
            case 'table':
                self.render_table(tag)
            case 'pre':
                self.render_code(tag)
            case 'ul' | 'ol':
                self.render_list(tag)
            case _:
                rendered = False
        
        # Iterate
        if not rendered:
            for t in tag.contents:
                if t.name is None: continue
                self.render_tag(t, level+1)

        self.stack.pop()

    def render(self):
        self.render_tag(self.soup)
        return self.doc


arg_parser = argparse.ArgumentParser(description='Generate a Docx file from one or more Markdown files')
arg_parser.add_argument('output', default=None, help='Output file')
arg_parser.add_argument('--files', default="*.md", help='Regex for Markdown files')
args = arg_parser.parse_args()

file_data = []
for part in sorted(glob.glob(args.files)):
    with open(part, 'r', encoding="utf-8") as f:
        file_data.append(f.read())

html = ''.join([mistletoe.markdown(data) for data in file_data])

soup = bs(html, 'html.parser')
html = soup.decode()

apply_html_style(soup)

# For testing
open('test.html','w').write(str(soup))

document = HtmlToDocx(soup).render()

document.save(args.output)
